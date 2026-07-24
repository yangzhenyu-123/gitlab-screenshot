"""Word 文档生成模块。

将截图按"章节 → 子节 → 截图"层级插入 Word 文档。每张截图独占一页
（可通过 Config.continuous 配置为连续排列）。
"""
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Inches

from gitlabshot.config import Config
from gitlabshot.preprocess import preprocess_image


@dataclass
class SubSection:
    """子节：对应分支名、tag 名或单个 commit。

    level 指定标题层级（2 或 3）；context 章节的说明子节用 level=2，
    各 commit 子节用 level=3。images 为该子节的截图路径列表。
    """

    title: str
    images: List[Path] = field(default_factory=list)
    level: int = 2


@dataclass
class Chapter:
    """章：对应 Heading 1。

    subsections 为该章下属的子节列表。常规章节（分支文件树、Tags 列表、
    Tag Commits）子节用 Heading 2；context 章节含一个说明性 Heading 2
    子节及若干 Heading 3 commit 子节（由 SubSection.level 控制）。
    """

    title: str
    subsections: List[SubSection] = field(default_factory=list)


@dataclass
class DocContent:
    """整份文档的内容：由若干章组成。"""

    chapters: List[Chapter] = field(default_factory=list)


def build_docx(content: DocContent, output_path: str, config: Config) -> None:
    """根据 content 生成 Word 文档，保存到 output_path。"""
    doc = Document()

    # 1. 设置所有 section 的四向边距
    for section in doc.sections:
        section.top_margin = Inches(config.margin_inches)
        section.bottom_margin = Inches(config.margin_inches)
        section.left_margin = Inches(config.margin_inches)
        section.right_margin = Inches(config.margin_inches)

    # 2. 用第一个 section 计算页面可用宽度（英寸）
    first_section = doc.sections[0]
    available_width = (
        first_section.page_width.inches
        - first_section.left_margin.inches
        - first_section.right_margin.inches
    )

    # 3. 遍历章节，逐章插入标题与截图
    is_first_chapter = True
    for chapter in content.chapters:
        # 空章节跳过：所有 subsections 的 images 均为空
        if not any(sub.images for sub in chapter.subsections):
            continue

        # 章节间分页：从第二章起，在 Heading 1 之前插入分页符
        if not is_first_chapter:
            doc.add_page_break()
        is_first_chapter = False

        doc.add_heading(chapter.title, level=1)

        is_first_subsection = True
        for subsection in chapter.subsections:
            # 跳过既无标题又无图片的空子节
            if not subsection.title and not subsection.images:
                continue
            # continuous 模式：同 subsection 内不分页，但子节间仍分页
            if config.continuous and not is_first_subsection:
                doc.add_page_break()
            is_first_subsection = False

            if subsection.title:
                doc.add_heading(subsection.title, level=subsection.level)

            for img in subsection.images:
                processed = preprocess_image(img, config)
                doc.add_picture(str(processed), width=Inches(available_width))
                # 非 continuous 模式：每张图后分页；continuous 模式同 subsection 内不分页
                if not config.continuous:
                    doc.add_page_break()

    doc.save(output_path)
